from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


SOURCE = Path('PANDUAN_PENGGUNA_VITAE_BERGAMBAR.docx')
OUTPUT = Path('PANDUAN_PENGGUNA_VITAE_BERGAMBAR_TERBARU.docx')


def replace_paragraph(paragraph, text, bold_prefix=None):
    """Replace text but retain the original paragraph style and alignment."""
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    if bold_prefix and text.startswith(bold_prefix):
        run.clear()
        first = run.add_text(bold_prefix)
        run.bold = True
        run.add_text(text[len(bold_prefix):])


def insert_paragraph_after(paragraph, text='', style=None):
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    paragraph._p.addnext(new_paragraph._p)
    if style:
        new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def add_numbered_steps_after(paragraph, steps):
    anchor = paragraph
    for step in steps:
        anchor = insert_paragraph_after(anchor, step, 'List Number')
    return anchor


doc = Document(SOURCE)
paragraphs = doc.paragraphs
by_text = {p.text.strip(): p for p in paragraphs if p.text.strip()}

# Make the first page immediately actionable for first-time users.
replace_paragraph(
    by_text['Vitae adalah aplikasi untuk membantu karyawan menyusun CV. Pengguna melengkapi data yang belum tersedia, menyimpan draft, melihat preview, lalu mengunduh CV dalam format PDF.'],
    'Vitae membantu karyawan membuat CV secara bertahap. Ikuti panduan ini dari persiapan akun sampai file PDF berhasil diunduh. Data dapat disimpan sebagai draft dan dilanjutkan kapan saja.'
)
anchor = by_text['2. Alur Penggunaan Singkat']
insert_paragraph_after(
    anchor,
    'Ikuti urutan berikut jika Anda baru pertama kali menggunakan Vitae. Jangan langsung mengunduh PDF sebelum data wajib dan preview diperiksa.',
    'Normal'
)

# Convert ambiguous overview lines into a clear start-to-finish procedure.
overview_replacements = {
    'Buka website Vitae melalui browser.': '1. Buka Google Chrome atau Microsoft Edge, lalu masukkan alamat website Vitae yang diberikan HR/admin.',
    'Jika belum punya akun, klik Daftar dan isi NIK, tanggal lahir, email, password, serta konfirmasi password.': '2. Jika belum memiliki akun, pilih Daftar. Isi NIK dan tanggal lahir sesuai data V-People, lalu gunakan email aktif serta password minimal 8 karakter.',
    'Buka email dan klik link verifikasi.': '3. Buka email yang didaftarkan dan klik link verifikasi dari Vitae. Cek Spam/Junk bila email belum terlihat.',
    'Masuk menggunakan email atau NIK dan password.': '4. Setelah verifikasi berhasil, masuk menggunakan email atau NIK dan password.',
    'Buka Dashboard, lalu klik Lengkapi CV atau Lanjut Isi Draft.': '5. Dari Dashboard, klik Lengkapi CV untuk memulai atau Lanjut Isi Draft untuk meneruskan data yang pernah disimpan.',
    'Isi form CV dari Step 1 sampai Step 7.': '6. Isi form berurutan dari Step 1 sampai Step 7. Simpan draft setiap selesai mengisi bagian penting.',
    'Klik Simpan Draft atau Simpan & Preview.': '7. Pilih Simpan Draft bila belum selesai; pilih Simpan & Preview bila data ingin diperiksa dalam tampilan CV.',
    'Periksa preview. Jika sudah benar, klik Download PDF.': '8. Periksa preview. Setelah semua data wajib lengkap dan tampilannya benar, klik Download PDF.',
}
for old, new in overview_replacements.items():
    replace_paragraph(by_text[old], new)

replacements = {
    'Buka alamat website Vitae dari browser.': 'Buka Google Chrome atau Microsoft Edge, masukkan alamat Vitae dari HR/admin, lalu tunggu halaman Masuk terbuka.',
    'Isi Email atau NIK.': 'Masukkan Email Login atau NIK yang telah terdaftar.',
    'Isi Password.': 'Masukkan password akun Anda.',
    'Klik Masuk.': 'Klik Masuk. Jika berhasil, Anda akan masuk ke Dashboard.',
    'Klik Daftar dari halaman Masuk.': 'Dari halaman Masuk, klik Daftar untuk membuka formulir pembuatan akun.',
    'Isi Tanggal Lahir sesuai KTP.': 'Isi Tanggal Lahir sesuai data yang tercatat di V-People.',
    'Klik Validasi dan Buat Akun.': 'Periksa kembali semua isian, lalu klik Validasi dan Buat Akun. Lanjutkan ke verifikasi email setelah akun dibuat.',
    'Buka email yang didaftarkan.': 'Buka kotak masuk email yang Anda gunakan saat mendaftar.',
    'Cari email verifikasi dari Vitae.': 'Cari email verifikasi dari Vitae; bila tidak ada di Inbox, periksa Spam/Junk.',
    'Klik link verifikasi di dalam email.': 'Klik link verifikasi terbaru di dalam email. Setelah berhasil, kembali ke Vitae untuk masuk.',
    'Lihat ringkasan data karyawan, seperti nama, jabatan, email, nomor HP, divisi, dan area kerja.': 'Periksa ringkasan data dan status kelengkapan CV. Data bertanda V-People berasal dari sistem sumber dan biasanya tidak dapat diubah di Vitae.',
    'Klik Lengkapi CV untuk membuka form pengisian.': 'Klik Lengkapi CV untuk mengisi CV dari awal, atau Lanjut Isi Draft bila sebelumnya sudah pernah menyimpan data.',
    'Klik Preview CV untuk melihat tampilan CV yang terakhir tersimpan.': 'Klik Preview CV untuk melihat data dari draft terakhir yang tersimpan.',
    'Klik Download PDF jika data wajib sudah lengkap.': 'Klik Download PDF hanya setelah sistem menyatakan data wajib lengkap dan preview sudah diperiksa.',
    'Lengkapi tempat lahir, jenis kelamin, status pernikahan, nomor HP, alamat, dan posisi jika diperlukan.': 'Lengkapi tempat lahir, jenis kelamin, status pernikahan, nomor HP, alamat lengkap, dan posisi bila masih kosong. Kolom ini diperlukan untuk menghasilkan PDF.',
    'Pilih wilayah berurutan: Provinsi, Kabupaten/Kota, Kecamatan, lalu Kelurahan/Desa.': 'Pilih wilayah secara berurutan: Provinsi, Kabupaten/Kota, Kecamatan, lalu Kelurahan/Desa. Pilihan berikutnya baru aktif setelah pilihan sebelumnya dipilih.',
    'Klik Berikutnya atau Simpan Draft.': 'Klik Simpan Draft sebelum berpindah bila ingin menyimpan perubahan, lalu klik Berikutnya untuk melanjutkan ke Step 2.',
    'Isi nama posisi atau jabatan.': 'Isi nama posisi/jabatan dan nama perusahaan secara lengkap.',
    'Pilih bulan mulai dan bulan selesai.': 'Pilih bulan mulai dan bulan selesai. Bila masih bekerja, centang Masih bekerja sampai sekarang.',
    'Tulis tanggung jawab dalam bentuk poin agar mudah dibaca.': 'Tulis tanggung jawab utama dalam poin singkat, misalnya pekerjaan yang dilakukan, alat/sistem yang digunakan, atau hasil kerja.',
    'Pilih jenjang pendidikan, misalnya SMA, SMK, D3, S1, S2.': 'Pilih jenjang pendidikan, isi nama institusi, jurusan, serta tahun lulus.',
    'Isi Keahlian Teknis sesuai kemampuan kerja, misalnya SAP, AutoCAD, Welding, Microsoft Excel, atau Maintenance.': 'Isi Keahlian Teknis sesuai kemampuan nyata, misalnya SAP, AutoCAD, Welding, Microsoft Excel, atau Maintenance.',
    'Isi Keahlian Non-teknis, misalnya komunikasi, kepemimpinan, kerja sama tim, atau problem solving.': 'Isi Keahlian Non-teknis yang relevan, misalnya komunikasi, kepemimpinan, kerja sama tim, atau problem solving.',
    'Pisahkan keahlian dengan koma atau baris baru.': 'Pisahkan setiap keahlian dengan koma atau baris baru, lalu periksa kembali agar tidak ada duplikasi.',
    'Isi nama sertifikasi atau pelatihan.': 'Isi nama sertifikasi/pelatihan, penerbit atau penyelenggara, jenis, dan tahun.',
    'Pilih masa berlaku atau centang seumur hidup/tanpa masa berlaku.': 'Isi masa berlaku bila ada. Jika tidak memiliki masa berlaku, centang Seumur hidup/tanpa masa berlaku.',
    'Isi Bahasa dan tingkat kemampuan jika ada.': 'Tambahkan Bahasa dan tingkat kemampuan bila relevan.',
    'Isi Proyek yang relevan dengan pekerjaan jika ada.': 'Tambahkan Proyek yang relevan beserta tahun pelaksanaannya bila ada.',
    'Isi Organisasi, jabatan/peran, tahun mulai, dan tahun selesai jika ada.': 'Tambahkan Organisasi beserta jabatan/peran dan rentang tahun bila ada.',
    'Tulis ringkasan profil maksimal 300 karakter.': 'Tulis ringkasan profil maksimal 300 karakter. Sebutkan bidang kerja, pengalaman, dan keahlian utama secara padat.',
    'Klik Generate jika ingin sistem membantu membuat ringkasan.': 'Jika memakai Generate, pastikan pengalaman dan keahlian telah diisi. Baca hasilnya dan sesuaikan dengan kondisi Anda.',
    'Klik Simpan Draft atau Simpan & Preview.': 'Klik Simpan Draft bila ingin melanjutkan nanti, atau Simpan & Preview untuk memeriksa hasil CV sebelum diunduh.',
    'Periksa tampilan akhir CV pada halaman Preview.': 'Periksa nama, kontak, pengalaman, pendidikan, keahlian, dan ejaan pada halaman Preview.',
    'Klik Edit CV atau Edit Draft jika ada data yang perlu diperbaiki.': 'Jika ada data yang salah atau belum tampil, klik Edit CV/Edit Draft, perbaiki data, lalu simpan kembali.',
    'Jika tampilan sudah benar dan data wajib lengkap, klik Download PDF.': 'Jika preview sudah benar dan data wajib lengkap, klik Download PDF. File CV akan diunduh oleh browser.',
}
for old, new in replacements.items():
    if old in by_text:
        replace_paragraph(by_text[old], new)

# The same original text also appears in the profile section; by_text points to
# the last occurrence, so handle the overview occurrence explicitly.
for paragraph in doc.paragraphs:
    if paragraph.text.strip() == 'Klik Simpan Draft atau Simpan & Preview.':
        replace_paragraph(
            paragraph,
            '7. Pilih Simpan Draft bila belum selesai; pilih Simpan & Preview bila data ingin diperiksa dalam tampilan CV.'
        )

# Add a usable pre-download checklist immediately before the troubleshooting section.
tips_heading = by_text['4. Tips Mengisi CV']
check_heading = insert_paragraph_after(tips_heading, 'Checklist Sebelum Download PDF', 'Heading 2')
anchor = add_numbered_steps_after(check_heading, [
    'Pastikan tempat lahir, jenis kelamin, status pernikahan, nomor HP, dan alamat lengkap sudah diisi.',
    'Pastikan ringkasan profil, keahlian teknis, minimal satu pengalaman kerja, dan minimal satu pendidikan sudah diisi.',
    'Klik Simpan Draft atau Simpan & Preview agar perubahan terakhir tersimpan.',
    'Buka Preview CV dan periksa kembali seluruh informasi sebelum mengunduh PDF.',
])

# Add explicit recovery guidance at the bottom, before the existing issue table.
issue_heading = by_text['5. Kendala Umum dan Solusi']
insert_paragraph_after(
    issue_heading,
    'Jika muncul pesan kesalahan, baca pesan tersebut terlebih dahulu. Jangan membagikan password kepada siapa pun. Saat menghubungi HR/admin, sertakan nama, NIK, halaman yang bermasalah, dan screenshot bila memungkinkan.',
    'Normal'
)

doc.core_properties.title = 'Panduan Pengguna Vitae'
doc.core_properties.subject = 'Panduan penggunaan Vitae dari pendaftaran sampai unduh PDF'
doc.save(OUTPUT)
print(OUTPUT)
